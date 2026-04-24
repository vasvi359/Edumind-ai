# rag.py

import os
from docx import Document as DocxDocument
import PyPDF2
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Global embedding model — loaded once, reused
_embedding = None


def get_embedding():
    global _embedding
    if _embedding is None:
        from langchain_huggingface import HuggingFaceEmbeddings
        _embedding = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return _embedding


def extract_text(uploaded_file):
    """
    Extracts text from PDF or Word document.
    Returns raw text string.
    """
    text = ""
    try:
        if uploaded_file.name.endswith(".docx"):
            doc = DocxDocument(uploaded_file)
            text = "\n".join([
                p.text for p in doc.paragraphs
                if p.text.strip()
            ])

        elif uploaded_file.name.endswith(".pdf"):
            pdf = PyPDF2.PdfReader(uploaded_file)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

    except Exception as e:
        print(f"Error extracting text: {e}")

    return text


def create_vectorstore(text):
    """
    Splits text into chunks, converts to embeddings,
    stores in ChromaDB. Returns (vectorstore, chunk_count).
    """
    from langchain_community.vectorstores import Chroma

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=700,
        chunk_overlap=100
    )

    chunks = splitter.split_text(text)
    docs = [Document(page_content=chunk) for chunk in chunks]

    vectorstore = Chroma.from_documents(
        documents=docs,
        embedding=get_embedding()
    )

    return vectorstore, len(chunks)


def retrieve_context(query, vectorstore):
    """
    Searches ChromaDB for most relevant chunks.
    Returns combined context string.
    """
    if vectorstore is None:
        return ""

    docs = vectorstore.similarity_search(query, k=4)
    return "\n".join([d.page_content for d in docs])
